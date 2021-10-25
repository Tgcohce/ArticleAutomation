import ArticleCmp
import ArticleInp
import docx

c = ArticleCmp
t = ArticleInp
doc = docx.Document()

temp11 = ("If you’re trying to decide on which sportsbook to go with, you may be considering " + t.FSNA + " and "
            + ArticleInp.FSNB + "." + " Both of these sportsbooks are legal and licensed in U.S. states but have unique offerings.")

temp12 = ("Our " + c.SSNA + " vs. " + c.SSNB + " sportsbook comparison will help you decide "
                                              "which is better for your sports betting needs.")
temp13 = (
"To start, we rated " + t.FSNA + " as " + c.RSA + " out of 5, or " + c.RTA + " and " + t.FSNB + " as " + c.RSB + " stars," +
" or " + c.RTB + ". This means if you’re just looking for the best overall sportsbook, then ")

temp14 = (t.FSNA + " is available in:")
temp15 = (t.FSNB + " is available in:")

temp21 = ( "When it comes to the intro promo for new customers,"
           " we rated " + t.FSNA + " as " + c.ipA + " stars and " + t.FSNB + " as " + c.ipB + " stars." )

temp22 = (t.FSNA + " offers a " + c.ipdescA + " that is refunded as " + c.ipcredA + "." + " You have " +
          c.ipcredtimeA + " to use the credit and there is a " + c.ipplayreqA + " playthrough requirement.")

temp23 = (t.FSNB + " offers a " + c.ipdescB + " that is refunded as " + c.ipcredB + "." + " You have " +
          c.ipcredtimeB + " to use the credit and there is a " + c.ipplayreqB + " playthrough requirement.")

temp30 = (t.FSNA + " vs " + t.FSNB + " Odds Competitiveness ")

temp31 = ("When it comes to the competitiveness of odds, we rated " + c.SSNA + " as " +
c.oddscompA + " out of 5" + " and " + c.SSNB + " as " + c.oddscompB + " out of 5.")

temp32 = ("To come up with these ratings, we analyzed hundreds of games in all major sports and leagues to find which books, "
          "on average, give customers the best deals.")

temp33 = ("If you’re looking for the most competitive odds, overall, then")

temp40 = (t.FSNA + " vs " + t.FSNB + " Promotions & Odds Boosts")

temp41 = ("When it comes to ongoing promotions and odds boosts that all customers can take advantage of, "
          "we rated " + t.FSNA + " as " + c.pobA + " stars and " + t.FSNB + " as " + c.pobB + " stars.")

temp42 = ("Odds boosts and promotions vary from sport to sport, "
          "so check with each to see which offers better specials on the sports you care about the most.")

temp50 = (t.FSNA + " vs " + t.FSNB + " User Interface " )

temp51 = ("You can typically bet with the " + c.SSNA + " and " + c.SSNB + " sportsbooks via their websites and"
          " mobile apps (iOS and Android). ")


temp52 = ("We rated the " + t.FSNA + " user interface as " +
          c.uiA + " stars, or " + c.uiRTA + " and " + t.FSNB + " as " + c.uiB + " stars, or " + c.uiRTB + ".")

temp53 = ("While most legitimate online sportsbooks in the U.S. are similarly designed and are easy enough to navigate, "
          "there are small things here and there that can make using one much more enjoyable than the other.")

temp54 = "Before registering, it may be worth downloading both apps or visiting their websites to see which you like better."


temp61 = ("Almost all major sportsbooks offer all of the main types of sports bets, including spreads,"
          " totals, moneyline bets, parlays, teasers, props, and futures, as well as live betting and cashing out of bets early.")

temp62 = ("One thing that not all offer, however, is same game parlays which allow you to combine sides, totals, and "
          "props from one game into a single parlay. ")

temp63 = ("As of this writing,")

temp71 = ("Having a variety of methods to deposit money into your account and to transfer your money out of it is another" 
         " important thing when it comes to finding the best sportsbook.")

temp72 = ("We rated " + t.FSNA + "'s " + "banking options as " + c.DepA + " and " + t.FSNB + "'s as " + c.DepB + " stars.")

temp73 = (t.FSNA + " offers " + c.DepmetA + " deposit methods and the minimum deposit amount is " +
          c.MindepA + ". It also has " + c.WimetA + " withdrawal methods and the minimum withdrawal is " + c.MinwitA + ".")

temp74 = (t.FSNB + " offers " + c.DepmetB + " deposit methods and the minimum deposit amount is " +
          c.MindepB + ". It also has " + c.WimetB + " withdrawal methods and the minimum withdrawal is " + c.MinwitB + ".")

temp80 = ("While our overall rating shows that ")

temp81 = (c.boline + " it’s important to think about what you value the most when it comes to sports betting. ")

temp82 = ("Use the comparison tool above to directly compare the things that are most valuable to you and check out our full reviews to learn more "
          "about all of the things we briefly touched on in this article. ")

temp83 = ("And remember, there is really no limit to the number of sportsbooks you can sign up for. Actually, it’s always a good idea to have at least a few that you can bet with so you can shop around for the best lines and odds, "
          "and so you can take advantage of the promotions that each offers.")
doc.add_paragraph("  ")
doc.add_paragraph("  ")
doc.add_paragraph(temp11)
doc.add_paragraph(temp12)
doc.add_paragraph(temp13 + '' + c.Ovr)

doc.add_paragraph(temp14)
for i in range(0, len(c.states_available)):
    doc.add_paragraph(c.states_available[i], style='List Bullet 3')

doc.add_paragraph(temp15)
for i in range(0, len(c.states_available_2)):
    doc.add_paragraph(c.states_available_2[i], style='List Bullet 3')

doc.add_heading(t.FSNA + " vs. " + t.FSNB + " Intro Promo ")
doc.add_paragraph(temp21)
doc.add_paragraph(temp22)
doc.add_paragraph(temp23)
doc.add_heading(temp30)
doc.add_paragraph(temp31)
doc.add_paragraph(temp32)
doc.add_paragraph(temp33 + ' ' + c.Ocmp)
doc.add_heading(temp40)
doc.add_paragraph(temp41 + " " + c.Pobdesc)
doc.add_paragraph(temp42)
doc.add_heading(temp50)
doc.add_paragraph(temp52)
doc.add_paragraph(temp53)
doc.add_paragraph(temp54)
doc.add_heading(t.FSNA + " vs. " + t.FSNB + " Bet Types")
doc.add_paragraph(temp61)
doc.add_paragraph(temp62)
doc.add_paragraph(temp63 + " " + c.Betdesc)
doc.add_heading(t.FSNA + " vs. " + t.FSNB + " Deposits & Withdrawals")
doc.add_paragraph(temp71)
doc.add_paragraph(temp72)
doc.add_paragraph(temp73)
doc.add_paragraph(temp74)
doc.add_heading("Bottom Line on if " + t.FSNA + " or " + t.FSNB + " is Better ")
doc.add_paragraph(temp80 + temp81)
doc.add_paragraph(temp82)
doc.add_paragraph(temp83)
doc.save('Article.docx')
